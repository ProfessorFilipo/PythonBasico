from docx import Document

# Cria um novo documento Word
doc = Document()

# Adiciona um título
doc.add_heading('Tabela de Comandos R Utilizados', level=1)

# Dados da tabela
dados = [
    ["Comando", "Função", "Exemplo"],
    ["readline(prompt = \"\")",
     "Lê entrada do usuário como string",
     "readline(prompt = \"Digite uma entrada\")"],
    ["as.numeric()",
     "Converte string para número",
     "as.numeric(\"123\")"],
    ["cat()",
     "Exibe mensagem formatada na tela",
     "cat(\"Texto\", variavel)"],
    ["print()",
     "Exibe valor de variáveis ou expressões",
     "print(variavel)"],
    ["length()",
     "Retorna o tamanho de vetor ou lista",
     "length(c(1,2,3))"],
    ["sort()",
     "Ordena elementos de um vetor",
     "sort(c(3,1,2))"],
    ["max()",
     "Retorna o maior valor",
     "max(c(1,5,3))"],
    ["min()",
     "Retorna o menor valor",
     "min(c(1,5,3))"],
    ["sum()",
     "Soma todos os elementos",
     "sum(c(1,2,3))"],
    ["/",
     "Operador de divisão",
     "5 / 2"],
    ["%%",
     "Resto da divisão",
     "7 %% 3"],
    ["4 %% 2 == 0",
     "Verifica se é par",
     "Verifica se resto da divisão por 2 é zero"],
    ["for (i in 1:5) { ... }",
     "Laço de repetição",
     "Executa alguma ação várias vezes"],
    ["sqrt()",
     "Calcula raiz quadrada",
     "sqrt(16)"],
    ["round()",
     "Arredonda número",
     "round(3.14159, 2)"],
    ["rev() / strsplit() + paste()",
     "Inverte elementos de vetor ou caracteres",
     "rev(c(1,2,3)) ou strsplit() + paste()"]
]

# Adiciona a tabela ao documento
table = doc.add_table(rows=1, cols=3)

# Adiciona cabeçalhos
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Comando'
hdr_cells[1].text = 'Função'
hdr_cells[2].text = 'Exemplo'

# Preenche as linhas
for linha in dados[1:]:
    row_cells = table.add_row().cells
    for i, valor in enumerate(linha):
        row_cells[i].text = valor

# Salva o arquivo
caminho_arquivo = 'Tabela_Comandos_R_Python.docx'
doc.save(caminho_arquivo)

print(f"Arquivo '{caminho_arquivo}' foi criado com sucesso!")
